import html
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def docx_to_html(path: str | Path) -> str:
    """将 .docx 文件转换为可用于在线预览的简化 HTML。

    仅输出段落和表格，避免执行脚本或外链资源。
    """
    doc = Document(str(path))
    paragraphs = {id(p._element): p for p in doc.paragraphs}
    tables = {id(t._element): t for t in doc.tables}

    parts = [
        "<!DOCTYPE html>",
        '<html><head><meta charset="utf-8"><style>',
        "body{font-family:sans-serif;line-height:1.6;padding:20px;max-width:800px;margin:0 auto;}",
        "p{margin:0 0 .5em 0;}",
        "table{border-collapse:collapse;width:100%;margin:1em 0;}",
        "td,th{border:1px solid #ccc;padding:6px;vertical-align:top;}",
        "th{background:#f5f5f5;}",
        "</style></head><body>",
    ]

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            para = paragraphs.get(id(child))
            if para is None:
                continue
            text = para.text.strip()
            if text:
                parts.append(f"<p>{html.escape(text)}</p>")
        elif child.tag == qn("w:tbl"):
            table = tables.get(id(child))
            if table is None:
                continue
            parts.append("<table>")
            for row in table.rows:
                parts.append("<tr>")
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    parts.append(f"<td>{html.escape(cell_text)}</td>")
                parts.append("</tr>")
            parts.append("</table>")

    parts.append("</body></html>")
    return "\n".join(parts)
